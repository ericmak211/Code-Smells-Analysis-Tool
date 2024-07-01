import subprocess
import sys
import os
import shutil
from datetime import datetime, timedelta
from git import Repo
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import difflib

records = ()

def clone_repository(repo_url, clone_dir):
    try:
        if os.path.exists(clone_dir):
            shutil.rmtree(clone_dir)
        Repo.clone_from(repo_url, clone_dir)
        print(f"Repository cloned to {clone_dir}")
        return True
    except Exception as e:
        print(f"An error occurred while cloning the repository: {e}")
        return False

def find_python_files(directory):
    python_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith('.py'):
                python_files.append(os.path.join(root, file))
    return python_files

def find_java_files(directory):
    java_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith('.java'):
                java_files.append(os.path.join(root, file))
    return java_files

def select_main_python_file(python_files):
    # Heuristic to prioritize certain files
    priority_files = ['main.py', 'app.py']
    for priority_file in priority_files:
        for file in python_files:
            if file.endswith(priority_file):
                return file
    # If no priority file is found, return the first Python file found
    if python_files:
        return python_files[0]
    return None

def select_main_java_file(java_files):
    priority_files = ['Main.java', 'App.java']
    for priority_file in priority_files:
        for file in java_files:
            if file.endswith(priority_file):
                return file
    if java_files:
        return java_files[0]
    return None

def calculate_code_similarity(old_code, new_code):
    diff = difflib.ndiff(old_code.splitlines(), new_code.splitlines())
    changes = [line for line in diff if line.startswith(('+', '-'))]
    return 1 - (len(changes) / max(len(old_code.splitlines()), len(new_code.splitlines())))

def check_refactoring_frequency(repo_path, file_path, days):
    try:
        repo = Repo(repo_path)
        # Convert backslashes to forward slashes for Git compatibility
        git_file_path = file_path.replace('\\', '/')
        commits = list(repo.iter_commits(paths=git_file_path, since=(datetime.now() - timedelta(days=days)).isoformat()))

        print(f"\nNumber of commits affecting {git_file_path} in the last {days} days: {len(commits)}")

        if len(commits) < 2:
            print("Not enough commits to analyze refactoring frequency.")
            return

        refactoring_time_ratios = []
        for i in range(1, len(commits)):
            commit_time = commits[i-1].committed_datetime
            previous_commit_time = commits[i].committed_datetime
            time_diff = (commit_time - previous_commit_time).total_seconds() / 3600  # in hours

            old_code = (commits[i].tree / git_file_path).data_stream.read().decode()
            new_code = (commits[i-1].tree / git_file_path).data_stream.read().decode()

            similarity = calculate_code_similarity(old_code, new_code)
            refactoring_time_ratio = (1 - similarity) * time_diff
            refactoring_time_ratios.append(refactoring_time_ratio)

        avg_refactoring_time_ratio = sum(refactoring_time_ratios) / len(refactoring_time_ratios)
        print(f"Average refactoring time ratio: {avg_refactoring_time_ratio:.2f}")

        add_report_line(git_file_path,len(commits), avg_refactoring_time_ratio)

        return avg_refactoring_time_ratio
    except Exception as e:
        print(f"An error occurred while analyzing the repository: {e}")
        return -1

def add_report_line(git_file_path, num_commits, avg_refactoring_time_ratio):
    global records

    y = ((git_file_path, str(num_commits), str(avg_refactoring_time_ratio)),)

    records += y

def generate_report(days):

    document = Document()

    ct = datetime.now().replace(second=0, microsecond=0)

    current_date = str(ct).replace(':', '').split(' ')[0]

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('Content Header', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(18)

    document.add_heading('Refactoring Activities Report for ' + current_date, 0)

    p = document.add_paragraph('')
    p.add_run('Executive Summary', style = 'Content Header').bold = True

    document.add_paragraph('This report provides an overview of the refactoring activities undertaken during the project timeline. The report highlights the commits and refactoring ratios for each file and offers recommendations for optimizing future refactoring efforts.')

    document.add_paragraph('Tools and Software Used:', style='List Bullet')
    document.add_paragraph('IDE Plugins: Refactoring Time Management Plugin', style='List Bullet 2')
    document.add_paragraph('Analysis Tools: SonarQube, CodeClimate', style='List Bullet 2')

    p = document.add_paragraph('')
    p.add_run('File Summary', style = 'Content Header').bold = True

    global records

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Class Name'
    hdr_cells[1].text = 'Commit Ct.'
    hdr_cells[2].text = 'Avg Refactor Time Ratio'
    for class_name, commit_count, refactor_ratio in records:
        row_cells = table.add_row().cells
        row_cells[0].text = class_name
        row_cells[1].text = commit_count
        row_cells[2].text = refactor_ratio

    p = document.add_paragraph('')
    p.add_run('Recommendations', style = 'Content Header').bold = True
    p1 = document.add_paragraph('Balance Refactoring and Development: Aim to Maintain refactoring activities at around 20%% of total development time to ensure new feature development remains on track.', style='List Number')
    p2 = document.add_paragraph('Prioritize High-Impact Refactoring: Focus on refactoring activities that significantly reduce code complexity and code smells', style='List Number')
    p3 = document.add_paragraph('Automate Code Quality Checks: Implement continuous integration tools to automatically check for code smells and complexity, ensuring early detection and resolution.', style='List Number')

    document.save('Refactoring Report - ' + str(ct).replace(':', '') + '.docx')

def main(repo_url, days):
    
    clone_dir = 'cloned_repo'

    # Check if clone directory already exists
    if not os.path.exists(clone_dir):
        # Clone the repository if it doesn't exist
        if not clone_repository(repo_url, clone_dir):
            return
    else:
        print(f"Using existing clone directory: {clone_dir}")

    # Find Python and Java files in the cloned repository
    python_files = find_python_files(clone_dir)
    java_files = find_java_files(clone_dir)

    # Combine Python and Java files for flexibility in analysis
    all_files = python_files + java_files

    if not all_files:
        print("No Python or Java files found in the repository.")
        return

    # Analyze refactoring frequency for each file
    for file_path in all_files:
        print(f"\nSelected file for analysis: {file_path}")

        # Check refactoring frequency
        file_path_in_repo = os.path.relpath(file_path, clone_dir)
        refactor_ratio = check_refactoring_frequency(clone_dir, file_path_in_repo, days)

        if refactor_ratio is None:
            print("Error occurred during refactoring frequency analysis.")
            continue
        elif refactor_ratio == -1:
            print("Not enough commits to analyze refactoring frequency.")
            continue

        # Provide insights based on refactoring frequency for the current file
        if refactor_ratio < 0.1:
            print("\nRefactoring is recommended based on the observed refactoring time ratios.")
        elif refactor_ratio > 0.5:
            print("\nIt seems like there has been frequent refactoring recently. Consider stabilizing the code.")
        else:
            print("\nThe code seems stable with minimal recent changes.")

    generate_report(days)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python check_refactoring.py <repository_url> <Number_Of_Days>")
    else:
        repo_url = sys.argv[1]
        try:
            days = int(sys.argv[2])
        except ValueError:
            print("NumberOfDays must be an integer.")
            sys.exit(1)
        main(repo_url, days)
